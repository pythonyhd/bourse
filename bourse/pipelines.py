# -*- coding: utf-8 -*-
import math
import os
import re
import time

import logging
from functools import reduce

import scrapy
from scrapy.exceptions import DropItem
from scrapy.pipelines.images import FilesPipeline

from bourse import settings
from bourse.settings import FILES_STORE
from bourse.utils.elastic_util import EsObject
from bourse.utils.filter_fact import cf_filter_fact

logger = logging.getLogger(__name__)


class BoursePipeline(object):
    """
    简单数据清洗
    """
    def process_item(self, item, spider):
        item['sj_ztxx'] = 1  # es显示数据专用
        # 添加分类标记
        item['sj_type'] = '15'
        # 添加时间戳
        item['cj_sj'] = math.ceil(time.time())
        ws_pc_id = cf_filter_fact(item)
        if ws_pc_id:
            item['ws_pc_id'] = ws_pc_id
        else:
            DropItem(item)
        # print(item)
        return item


class DownloadFilesPipeline(FilesPipeline):
    """
    文件下载，管道文件
    """

    def get_media_requests(self, item, info):
        file_url = item.get('xq_url', '')
        if file_url.endswith('pdf'):
            yield scrapy.Request(file_url)

        elif file_url.endswith('doc'):
            yield scrapy.Request(file_url)

        elif file_url.endswith('docx'):
            yield scrapy.Request(file_url)

        else:
            logger.debug('既不是PDF也不是word纯文本')

    def item_completed(self, results, item, info):
        re_com = re.compile(r'\r|\n|\t|\s')
        image_paths = [x['path'] for ok, x in results if ok]
        if not image_paths:
            item['cf_file_name'] = ''
            item['wbbz'] = '纯HTML格式'
        else:
            file_path = os.path.abspath(FILES_STORE) + r'/' + image_paths[0]
            item['cf_file_name'] = image_paths[0]
            if image_paths[0].endswith('doc'):
                docx_list = self.parse_doc2docx(file_path)
                if docx_list:
                    docx_text = reduce(lambda x, y: x + y, [re_com.sub('', i) for i in docx_list])
                    oname_pattern = re.compile(r'(关于对)(.*?公司)')
                    cf_sy_pattern = re.compile(r'(存在以下问题：|违规事实：|存在以下违规行为：|存在以下违规事实：)(.*?。)')
                    cf_yj_pattern = re.compile(r'((违反了本所|依据本所|根据).*?规定)')
                    cf_jg_pattern = re.compile(r'((本所决定|本所作出如下处分：|本所作出如下处分决定：).*?。)')
                    oname = oname_pattern.search(docx_text)
                    item['oname'] = oname.group(2) if oname else ''
                    cf_sy = cf_sy_pattern.search(docx_text)
                    item['cf_sy'] = cf_sy.group(2) if cf_sy else ''
                    cf_yj = cf_yj_pattern.search(docx_text)
                    item['cf_yj'] = cf_yj.group(1) if cf_yj else ''
                    cf_jg = cf_jg_pattern.search(docx_text)
                    item['cf_jg'] = cf_jg.group(1) if cf_jg else ''
                    item['ws_nr_txt'] = docx_text
                else:
                    logger.info('深圳证券交易所-获取不到word文档里面的内容--只把基本内容存到数据库')
            elif image_paths[0].endswith('docx'):
                docx_list = self.parse_docx(file_path)
                if docx_list:
                    docx_text = reduce(lambda x, y: x + y, [re_com.sub('', i) for i in docx_list])
                    oname_pattern = re.compile(r'(关于对)(.*?公司)')
                    cf_sy_pattern = re.compile(r'(存在以下问题：|违规事实：|存在以下违规行为：|董事会 ：|存在以下违规事实：)(.*?。)')
                    cf_yj_pattern = re.compile(r'((违反了本所|依据本所|根据|是否符合).*?规定)')
                    cf_jg_pattern = re.compile(r'((本所决定|本所作出如下处分：|本所作出如下处分决定：|请你公司).*?。)')
                    oname = oname_pattern.search(docx_text)
                    item['oname'] = oname.group(2) if oname else ''
                    cf_sy = cf_sy_pattern.search(docx_text)
                    item['cf_sy'] = cf_sy.group(2) if cf_sy else ''
                    cf_yj = cf_yj_pattern.search(docx_text)
                    item['cf_yj'] = cf_yj.group(1) if cf_yj else ''
                    cf_jg = cf_jg_pattern.search(docx_text)
                    item['cf_jg'] = cf_jg.group(1) if cf_jg else ''
                    item['ws_nr_txt'] = docx_text
                else:
                    logger.info('直接读取docx文件失败')
                return item

            else:
                logger.debug('不是word文件')

        # print(item)
        return item

    def parse_doc2docx(self, path):
        """
        解析docx文件内容，如果是doc文件需要先转换成docx文件格式
        :param response:
        :return:
        """
        import os
        from win32com.client import Dispatch
        from docx import Document
        try:
            word = Dispatch('Word.Application')
            word.Visible = 0
            word.DisplayAlerts = 0
            doc = word.Documents.Open(path)
            newpath = os.path.splitext(path)[0] + '.docx'
            doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
            doc.Close()
            word.Quit()
            os.remove(path)
            time.sleep(0.01)
            doc_text = Document(newpath)
            docText = [paragraph.text for paragraph in doc_text.paragraphs]
            return docText
        except Exception as e:
            print(f'转换出错:{repr(e)}')

    def parse_docx(self, docx_path):
        """
        直接读取docx
        :param path:
        :return:
        """
        from docx import Document
        # 打开文档
        document = Document(docx_path)
        # 处理纯文本
        docx_text = [paragraph.text for paragraph in document.paragraphs]
        return docx_text


class DealShanghaiFilesPipeline(FilesPipeline):
    """ 上海证券交易所-文件下载-doc-docx解析 """
    def get_media_requests(self, item, info):
        file_url = item.get('xq_url', '')
        if file_url.endswith('pdf'):
            yield scrapy.Request(file_url)

        elif file_url.endswith('doc'):
            yield scrapy.Request(file_url)

        elif file_url.endswith('docx'):
            yield scrapy.Request(file_url)

        else:
            logger.debug('既不是PDF也不是word纯文本')

    def item_completed(self, results, item, info):
        re_com = re.compile(r'\r|\n|\t|\s')
        image_paths = [x['path'] for ok, x in results if ok]
        if not image_paths:
            item['cf_file_name'] = ''
            item['wbbz'] = '纯HTML格式'
        else:
            file_path = os.path.abspath(FILES_STORE) + r'/' + image_paths[0]
            item['cf_file_name'] = image_paths[0]
            if image_paths[0].endswith('doc'):
                docx_list = self.parse_doc2docx(file_path)
                if docx_list:
                    cf_cfmc = docx_list[0].replace('标题：', '')
                    cf_sy = docx_list[-1].replace('处理事由：', '')
                    docx_text = reduce(lambda x, y: x + y, [re_com.sub('', i) for i in docx_list])
                    cf_sy_second = re.search(r'(经查明|经审核|经查.*?。)', docx_text)
                    cf_sy_second = cf_sy_second.group(1) if cf_sy_second else ''
                    oname_one = re.search(r'(关于对|当事人：|关于)(.*?)(名下证券账户)', docx_text)
                    oname_one = oname_one.group(2) if oname_one else ''
                    oname_two = re.search(r'(关于对|当事人：|关于)(.*?公司)', docx_text)
                    oname_two = oname_two.group(2) if oname_two else ''
                    oname = oname_one if oname_one else oname_two
                    item['oname'] = oname if oname else item.get('oname')
                    item['cf_cfmc'] = cf_cfmc if cf_cfmc else ''
                    cf_sy = cf_sy if cf_sy else cf_sy_second
                    item['cf_sy'] = cf_sy_second if '年' in cf_sy else cf_sy
                    item['ws_nr_txt'] = docx_text if docx_text else ''
                    cf_yj = re.search(r'(根据.*?规定)', docx_text)
                    cf_yj = cf_yj.group(1) if cf_yj else ''
                    item['cf_yj'] = cf_yj
                    cf_jg = re.search(r'((做出如下纪律处分决定：|本所决定|决定对).*?。)', docx_text)
                    cf_jg = cf_jg.group(1) if cf_jg else ''
                    item['cf_jg'] = cf_jg
                else:
                    logger.info("上海证券交易所-读取不到word文件")

            elif image_paths[0].endswith('docx'):
                logger.info('是docx文件，不需要转换')

            else:
                logger.debug('不是word文件')
        return item

    def parse_doc2docx(self, path):
        """ 解析docx文件内容，如果是doc文件需要先转换成docx文件格式 """
        import os
        from win32com.client import Dispatch
        from docx import Document
        try:
            word = Dispatch('Word.Application')
            word.Visible = 0
            word.DisplayAlerts = 0
            doc = word.Documents.Open(path)
            newpath = os.path.splitext(path)[0] + '.docx'
            doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
            doc.Close()
            word.Quit()
            os.remove(path)
            time.sleep(0.01)
            doc_text = Document(newpath)
            docText = [paragraph.text for paragraph in doc_text.paragraphs]
            return docText
        except Exception as e:
            print(f'转换出错:{repr(e)}')

    def parse_docx(self, docx_path):
        """ 直接读取docx """
        from docx import Document
        # 打开文档
        document = Document(docx_path)
        # 处理纯文本
        docx_text = [paragraph.text for paragraph in document.paragraphs]
        return docx_text


class DealNationalFilesPipeline(FilesPipeline):
    """ 全国中小企业股份转让系统-监管公开信息-问询函DOC跟DOCX """
    def get_media_requests(self, item, info):
        file_url = item.get('xq_url', '')
        if file_url.endswith('pdf'):
            yield scrapy.Request(file_url)

        elif file_url.endswith('doc'):
            yield scrapy.Request(file_url)

        elif file_url.endswith('docx'):
            yield scrapy.Request(file_url)

        else:
            logger.debug('既不是PDF也不是word纯文本')

    def item_completed(self, results, item, info):
        re_com = re.compile(r'\r|\n|\t|\s')
        image_paths = [x['path'] for ok, x in results if ok]
        if not image_paths:
            item['cf_file_name'] = ''
            item['wbbz'] = '纯HTML格式'
        else:
            file_path = os.path.abspath(FILES_STORE) + r'/' + image_paths[0]
            item['cf_file_name'] = image_paths[0]
            if image_paths[0].endswith('doc'):
                docx_list = self.parse_doc2docx(file_path)
                if docx_list:
                    docx_text = reduce(lambda x, y: x + y, [re_com.sub('', i) for i in docx_list])
                    cf_wsh_pattern = re.compile(r'(问询函)(半年报问询函.*?号)')
                    cf_sy_pattern = re.compile(r'((关注到以下情况：|经查明|经审阅|请你公司补充披露以下事项：).*?。)')
                    cf_yj_pattern = re.compile(r'((公司上述行为违反了|你的上述行为违反了|根据|依据).*?规定)')
                    cf_jg_pattern = re.compile(r'((做出如下纪律处分决定：|请就上述问题做出书面说明|收到本问询函后).*?。)')

                    cf_wsh = cf_wsh_pattern.search(docx_text)
                    cf_wsh = cf_wsh.group(2) if cf_wsh else ''
                    cf_sy = cf_sy_pattern.search(docx_text)
                    cf_sy = cf_sy.group(1) if cf_sy else ''
                    cf_yj = cf_yj_pattern.search(docx_text)
                    cf_yj = cf_yj.group(1) if cf_yj else ''
                    cf_jg = cf_jg_pattern.search(docx_text)
                    cf_jg = cf_jg.group(1) if cf_jg else ''
                    item['cf_wsh'] = cf_wsh
                    item['cf_sy'] = cf_sy
                    item['cf_yj'] = cf_yj
                    item['cf_jg'] = cf_jg
                else:
                    logger.info("全国中小企业股份转让系统-doc转换成docx出错-读取不到word文件")

            elif image_paths[0].endswith('docx'):
                docx_list = self.parse_doc2docx(file_path)
                if docx_list:
                    docx_text = reduce(lambda x, y: x + y, [re_com.sub('', i) for i in docx_list])
                    cf_wsh_pattern = re.compile(r'(问询函)(半年报问询函.*?号)')
                    cf_sy_pattern = re.compile(r'((关注到以下情况：|经查明|经审阅|请你公司补充披露以下事项：).*?。)')
                    cf_yj_pattern = re.compile(r'((公司上述行为违反了|你的上述行为违反了|根据|依据).*?规定)')
                    cf_jg_pattern = re.compile(r'((做出如下纪律处分决定：|请就上述问题做出书面说明|收到本问询函后).*?。)')

                    cf_wsh = cf_wsh_pattern.search(docx_text)
                    cf_wsh = cf_wsh.group(2) if cf_wsh else ''
                    cf_sy = cf_sy_pattern.search(docx_text)
                    cf_sy = cf_sy.group(1) if cf_sy else ''
                    cf_yj = cf_yj_pattern.search(docx_text)
                    cf_yj = cf_yj.group(1) if cf_yj else ''
                    cf_jg = cf_jg_pattern.search(docx_text)
                    cf_jg = cf_jg.group(1) if cf_jg else ''
                    item['cf_wsh'] = cf_wsh
                    item['cf_sy'] = cf_sy
                    item['cf_yj'] = cf_yj
                    item['cf_jg'] = cf_jg
                else:
                    logger.info("全国中小企业股份转让系统-读取不到word文件")

            else:
                logger.debug('不是word文件')

        # print(item)
        return item

    def parse_doc2docx(self, path):
        """ 解析docx文件内容，如果是doc文件需要先转换成docx文件格式 """
        import os
        from win32com.client import Dispatch
        from docx import Document
        try:
            word = Dispatch('Word.Application')
            word.Visible = 0
            word.DisplayAlerts = 0
            doc = word.Documents.Open(path)
            newpath = os.path.splitext(path)[0] + '.docx'
            doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
            doc.Close()
            word.Quit()
            os.remove(path)
            time.sleep(0.01)
            doc_text = Document(newpath)
            docText = [paragraph.text for paragraph in doc_text.paragraphs]
            return docText
        except Exception as e:
            print(f'转换出错:{repr(e)}')

    def parse_docx(self, docx_path):
        """ 直接读取docx """
        from docx import Document
        # 打开文档
        document = Document(docx_path)
        # 处理纯文本
        docx_text = [paragraph.text for paragraph in document.paragraphs]
        return docx_text


class Save2eEsPipeline(object):
    """ 存储elasticsearch """
    def __init__(self):
        self.es = EsObject(index_name=settings.INDEX_NAME, index_type=settings.INDEX_TYPE, host=settings.ES_HOST, port=settings.ES_PORT)

    def process_item(self, item, spider):
        if item:
            # 获取唯一ID
            _id = item['ws_pc_id']
            res1 = self.es.get_data_by_id(_id)
            if res1.get('found') == True:
                logger.debug("该数据已存在%s" % _id)
                # self.es.update_data(dict(item), _id)
            else:
                self.es.insert_data(dict(item), _id)
                logger.debug("----------抓取成功,开始插入数据%s" % _id)
                return item